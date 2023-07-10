from selenium import webdriver
from docx import Document

driver = webdriver.Chrome()
driver.get("")
screenshot_path=""
driver.save_screenshot(screenshot_path)

driver.quit()

#create a new word document
doc = Document()

doc.add_paragraph("Screenshot:")
doc.add_picture(screenshot_path)

#save the word Document
doc.save('screenshot.docx')

